from __future__ import annotations
from __future__ import annotations

import datetime
import io
import json
import os
import time
from http.client import OK

import docx
import openai
import pytz
import requests
from googleapiclient.discovery import build as google_build
from googleapiclient.errors import HttpError
from httplib2 import Http
from llama_index.core import ServiceContext, SimpleDirectoryReader, \
    Document, Settings
from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.vector_stores.weaviate import WeaviateVectorStore
from oauth2client.service_account import ServiceAccountCredentials
from pypdf import PdfReader

import weaviate
from business_units.models import BusinessUnit, TestUser
from eddy_school.settings import SEND_PULSE_URL, SMART_SENDER_URL
from pytorch_faq.utils import find_closest_answer

SEND_PULSE_AUTH = '/oauth/access_token'

SEND_PULSE_TELEGRAM_MESSAGE = '/telegram/contacts/sendText'
SEND_PULSE_TELEGRAM_RUN_BY_TRIGGER = '/telegram/flows/runByTrigger'

SEND_PULSE_VIBER_MESSAGE = '/viber/chatbots/contacts/send'
SEND_PULSE_VIBER_RUN_BY_TRIGGER = '/viber/chatbots/flows/runByTrigger'

SEND_PULSE_LIVE_CHAT_MESSAGE = '/live-chat/contacts/send'
SEND_PULSE_LIVE_CHAT_RUN_BY_TRIGGER = '/live-chat/flows/runByTrigger'

SMART_SENDER_MESSAGE = '/v1/contacts/{contactId}/send'
SMART_SENDER_RUN_BY_TRIGGER = '/v1/contacts/{contactId}/fire'

SCOPES = ['https://www.googleapis.com/auth/documents.readonly', 'https://www.googleapis.com/auth/drive']
DISCOVERY_DOC = 'https://docs.googleapis.com/$discovery/rest?version=v1'
DISCOVERY_DRIVE = 'https://www.googleapis.com/discovery/v1/apis/drive/v3/rest'

REFINE_PROMPT = (
    """
    ## User Query
        {query}



    ## Context
        {reference_answer}



    ## Generated Answer
        {generated_answer}
    """
)


def get_credentials(file_url):
    credentials = ServiceAccountCredentials.from_json_keyfile_name(
        file_url, SCOPES
    )
    return credentials


def read_paragraph_element(element):
    text_run = element.get('textRun')
    if not text_run:
        return ''
    return text_run.get('content')


def read_structural_elements(elements):
    text = ''
    for value in elements:
        if 'paragraph' in value:
            elements = value.get('paragraph').get('elements')
            for elem in elements:
                text += read_paragraph_element(elem)
        elif 'table' in value:
            table = value.get('table')
            for row in table.get('tableRows'):
                cells = row.get('tableCells')
                for cell in cells:
                    text += read_structural_elements(cell.get('content'))
        elif 'tableOfContents' in value:
            toc = value.get('tableOfContents')
            text += read_structural_elements(toc.get('content'))
    return text


def run_correctness_eval(
        query_str: str,
        reference_answer: str,
        generated_answer: str,
        llm: OpenAI,
        threshold: float = 0.0,
        eval_chat_template=None
):
    """Run correctness eval."""
    fmt_messages = eval_chat_template.format_messages(
        llm=llm,
        query=query_str,
        reference_answer=reference_answer,
        generated_answer=generated_answer,
    )
    chat_response = llm.chat(fmt_messages)
    raw_output = chat_response.message.content

    # Extract from response
    try:
        score_str, reasoning_str = raw_output.split("\n", 1)
        score = float(score_str)
        reasoning = reasoning_str.lstrip("\n")

        return {"passing": score >= threshold, "score": score, "reason": reasoning}
    except:
        try:
            score = float(raw_output)
            reasoning = reasoning_str.lstrip("\n")

            return {"passing": score >= threshold, "score": score, "reason": reasoning}
        except:
            score = raw_output.split(".")[0]

            return {"score": score}


def make_query(query_text, google_docs_ids, uploaded_files, documents_folder, index_name, openai_key, file_url,
               test_doc=False, contact_id=None):
    openai.api_key = openai_key
    os.environ["OPENAI_API_KEY"] = openai.api_key
    credentials = get_credentials(file_url)
    http = credentials.authorize(Http())
    business_unit = BusinessUnit.objects.filter(apikey=documents_folder.split('documents-')[1]).first()

    if test_doc:
        test_user, _ = TestUser.objects.get_or_create(contact_id=contact_id)

    closest_answer = find_closest_answer(business_unit.id, query_text, business_unit.similarity_simple_q)
    if closest_answer:
        return {"response": closest_answer, "eval_result": 5, "llm_context": 'None'}

    if business_unit.max_tokens:
        llm = OpenAI(model=business_unit.gpt_model, temperature=business_unit.temperature,
                     max_tokens=business_unit.max_tokens)
    else:
        llm = OpenAI(model=business_unit.gpt_model, temperature=business_unit.temperature)

    Settings.llm = llm
    Settings.embed_model = OpenAIEmbedding(model='text-embedding-3-large')

    f_doc_title = f'{business_unit.apikey}'
    docs_content = ""
    uploaded_files_ids = [document.id for document in uploaded_files]

    docs_list = google_docs_ids + uploaded_files_ids
    documents_changed = False

    if not os.path.exists(index_name):
        documents_changed = True

    if not business_unit.last_used_documents_list or business_unit.last_used_documents_list == '[]':
        documents_changed = True
        business_unit.last_used_documents_list = docs_list
        business_unit.save()
    elif business_unit.last_used_documents_list is [] or docs_list != eval(business_unit.last_used_documents_list):
        documents_changed = True
        business_unit.last_used_documents_list = docs_list
        business_unit.save()

    drive = google_build("drive", "v3", http=http)

    ukraine_timezone = pytz.timezone('Europe/Kiev')

    last_modified = business_unit.last_update_document.astimezone(
        ukraine_timezone) if business_unit.last_update_document else datetime.datetime.min.replace(
        tzinfo=pytz.utc).astimezone(ukraine_timezone)

    for document_id in google_docs_ids:
        try:
            modified_time = drive.files().get(fileId=document_id, fields='modifiedTime').execute()['modifiedTime']
        except HttpError as e:
            return {
                "response": f"{business_unit.panic_text if business_unit.panic_text else 'Oops. The document was not found or cannot be accessed.'}",
                "eval_result": 0,
                "llm_context": None}
        modified_datetime = datetime.datetime.strptime(modified_time, '%Y-%m-%dT%H:%M:%S.%fZ').replace(
            tzinfo=pytz.utc).astimezone(ukraine_timezone)
        if modified_datetime > last_modified:
            last_modified = modified_datetime
            documents_changed = True
            business_unit.last_update_document = modified_datetime.astimezone(pytz.utc)
            business_unit.save()

    if documents_changed or test_doc:
        docs_service = google_build('docs', 'v1', http=http)
        docs = []

        for document_id in google_docs_ids:
            try:
                doc = docs_service.documents().get(documentId=document_id).execute()
                docs_content += read_structural_elements(doc.get('body').get('content')) + "\n\n"
                docs.append(doc)
            except HttpError as e:
                return {"response": f"Failed to process Google Doc {document_id}: {str(e)}", "eval_result": 0,
                        "llm_context": None}

        for document in uploaded_files:
            if str(document.file).endswith('.docx'):
                doc = docx.Document(document.file.path)
                for paragraph in doc.paragraphs:
                    docs_content += paragraph.text + "\n"
            elif str(document.file).endswith('.pdf'):
                with open(document.file.path, 'rb') as f:
                    reader = PdfReader(f)
                    for page in reader.pages:
                        docs_content += page.extract_text() + "\n\n"
            else:
                try:
                    with document.file.open('rb') as f:
                        docs_content += f.read().decode('utf-8') + "\n\n"
                except:
                    pass

        filename = os.path.join(documents_folder, f_doc_title + '.txt')
        if not os.path.exists(documents_folder):
            os.makedirs(documents_folder, exist_ok=True)

        with io.open(filename, 'w', encoding='utf-8') as f:
            f.write(docs_content)

        business_unit.last_update_document = last_modified
        business_unit.save()

    temperature = business_unit.temperature
    if business_unit.max_tokens:
        llm = OpenAI(model=business_unit.gpt_model, temperature=temperature,
                     max_tokens=business_unit.max_tokens)
    else:
        llm = OpenAI(model=business_unit.gpt_model, temperature=temperature)

    service_context = ServiceContext.from_defaults(
        llm=llm,
        system_prompt=business_unit.system_prompt,
        chunk_size=business_unit.chunk_size if business_unit.chunk_size else None,
        chunk_overlap=business_unit.chunk_overlap if business_unit.chunk_overlap else None
    )

    documents = SimpleDirectoryReader(documents_folder).load_data()
    # auth_config = weaviate.AuthApiKey(api_key="f7myJDmYyg7q2CMTJN1vnQf1D3LaAE7d1ETj")

    client = weaviate.Client(
        url="http://sc.aiadmin.info:8080/",
        # auth_client_secret=auth_config
    )

    exists_test_index = False

    if test_doc:
        if test_user.file_hash_sum != documents[0].hash:
            test_user.file_hash_sum = documents[0].hash
            test_user.save()

            try:
                client.schema.delete_class(f'TestUser{contact_id}')
                print(f"Class TestUser{contact_id} has been deleted successfully.")
            except Exception as e:
                print(f"Failed to delete class TestUser{contact_id}: {str(e)}")
        else:
            exists_test_index = True

        vector_store = WeaviateVectorStore(
            weaviate_client=client,
            index_name=f"TestUser{contact_id}", text_key='content',
        )
    else:
        vector_store = WeaviateVectorStore(
            weaviate_client=client,
            index_name=f"Bu{business_unit.id}", text_key='content',
        )

    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    if os.path.exists(index_name) and documents_changed or (test_doc and not exists_test_index):
        VectorStoreIndex.from_documents(documents,
                                        storage_context=storage_context,
                                        show_progress=True
                                        )
    if not os.path.exists(index_name) and not test_doc:
        os.mkdir(index_name)
        VectorStoreIndex.from_documents(documents,
                                        storage_context=storage_context,
                                        show_progress=True
                                        )
    if os.path.exists(index_name) or exists_test_index:
        index = VectorStoreIndex.from_vector_store(vector_store, embed_model=Settings.embed_model,
                                                   service_context=service_context)

    query_engine = index.as_query_engine(
        similarity_top_k=business_unit.similarity_top_k if business_unit.similarity_top_k else 1
    )
    response = query_engine.query(query_text)

    context = []

    for context_info in response.source_nodes:
        context.append(Document(text=context_info.node.get_content()).text)
    return {"response": response.response, "eval_result": 5,
            "llm_context": context}


def translate_to_ukrainian(text):
    return text
    # translator = Translator()
    #
    # try:
    #     source_lang = translator.detect(text).lang
    #
    #     if source_lang == 'en':
    #         translation = translator.translate(text, src='en', dest='uk').text
    #         return translation
    #     else:
    #         return text
    # except:
    #     return text


def smart_sender_flow(request_type, business_units, contact_id=None, response=None):
    headers = {"Authorization": f"Bearer {business_units.smart_sender_token}"}

    if request_type == "send_message":
        r = requests.post(f'{SMART_SENDER_URL}{SMART_SENDER_MESSAGE.format(contactId=contact_id)}', headers=headers,
                          data={
                              "content": response,
                              "type": "text",
                              "watermark": int(time.time())
                          })
        return r
    elif request_type == "word_trigger":
        r = requests.post(f'{SMART_SENDER_URL}{SMART_SENDER_RUN_BY_TRIGGER.format(contactId=contact_id)}',
                          headers=headers,
                          data={
                              "name": business_units.default_text
                          })
        return r


def send_pulse_flow(request_type, business_units, contact_id=None, source_type=None, **kwargs):
    headers = {"Authorization": f"Bearer {business_units.sendpulse_token}"}

    if request_type == "auth":
        datetime_now = datetime.datetime.now(datetime.timezone.utc)

        data = {
            "grant_type": "client_credentials",
            "client_id": business_units.sendpulse_id,
            "client_secret": business_units.sendpulse_secret
        }
        sendpulse_auth = json.loads(requests.post(f'{SEND_PULSE_URL}{SEND_PULSE_AUTH}', data=data).text)

        business_units.sendpulse_token = sendpulse_auth.get('access_token', None)
        business_units.last_update_sendpulse = datetime_now
        business_units.save()
        return OK

    elif request_type == "word_trigger":
        send_pulse_url = {
            "telegram": SEND_PULSE_TELEGRAM_RUN_BY_TRIGGER,
            "viber": SEND_PULSE_VIBER_RUN_BY_TRIGGER,
            "live_chat": SEND_PULSE_LIVE_CHAT_RUN_BY_TRIGGER,
        }[source_type]

        request_data = {
            "contact_id": contact_id,
            "trigger_keyword": business_units.default_text,
        }

        r = requests.post(
            f"{SEND_PULSE_URL}{send_pulse_url}", headers=headers, data=request_data
        )

        return r

    elif request_type == "send_message":
        part = kwargs.get("part", None)

        if part and source_type == "telegram":
            r = requests.post(f'{SEND_PULSE_URL}{SEND_PULSE_TELEGRAM_MESSAGE}', headers=headers, data={
                "contact_id": contact_id,
                "text": part
            })
            return r

        elif part and source_type == "viber":
            r = requests.post(f'{SEND_PULSE_URL}{SEND_PULSE_VIBER_MESSAGE}', headers=headers, json={
                "contact_id": contact_id,
                "messages": [
                    {
                        "type": "text",
                        "text": {
                            "text": part
                        }
                    }
                ]
            })
            return r

        elif part and source_type == "live_chat":
            r = requests.post(f'{SEND_PULSE_URL}{SEND_PULSE_LIVE_CHAT_MESSAGE}', headers=headers, json={
                "contact_id": contact_id,
                "messages": [
                    {
                        "type": "text",
                        "text": {
                            "text": part
                        }
                    }
                ]
            })
            return r


def gpt_assistant_query(query_text, business_unit, openai_key):
    openai.api_key = openai_key
    os.environ["OPENAI_API_KEY"] = openai.api_key

    client = openai.OpenAI()
    thread = client.beta.threads.create()

    message = client.beta.threads.messages.create(
        thread_id=thread.id,
        role="user",
        content=query_text
    )

    run = client.beta.threads.runs.create(
        thread_id=thread.id,
        assistant_id=business_unit.gpt_assistant_id
    )

    while True:
        time.sleep(5)

        run_status = client.beta.threads.runs.retrieve(
            thread_id=thread.id,
            run_id=run.id
        )

        if run_status.status == 'completed':
            messages = client.beta.threads.messages.list(
                thread_id=thread.id
            )

            for msg in messages.data:
                role = msg.role
                content = msg.content[0].text.value
                if role.capitalize() == 'Assistant':
                    return {"response": content, "eval_result": 5,
                            "llm_context": "None, it's GPT assistant mode!"}

            break


def split_text_into_parts(text, max_length=512):
    final_parts = []
    while text:
        if len(text) <= max_length:
            final_parts.append(text)
            break

        split_index = text[:max_length].rfind('.') + 1
        if split_index <= 0:
            split_index = text[:max_length].rfind(' ') + 1
            if split_index <= 0:
                split_index = max_length

        final_parts.append(text[:split_index].strip())
        text = text[split_index:].strip()

    return final_parts
